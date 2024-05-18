# -*- coding: utf-8 -*-
"""
Created somewhere between March and May 2024

@author: Buro89

Case:
See docstring dobbelsteen.py

After playing the desired number of rounds in dobbelsteen.py, the scores are summarized with this script here: in an Excel file, 
but also in graphs and even a report!
The excel file, visualisations and report dynamically adjust to the number of rounds and players chosen by USER.

The plots:
- Graph 1 - (WIP) what is the total score for each participant?
- Graph 2 - (WIP) what is the average score for each participant?
- Graph 3 - (WIP) what is the total number of rounds for each participant that they had the top score?
- Graph 4 - (WIP) what is the number of top scores gained within each team, divided by the team size (no. of team members)?
- Graph 5 - (WIP) what is the amount of scores within the teams that are below the average total score? And equal & above?
- Graph 6 - (WIP) what is the average score for each team in Round 1 compared to Round 2? 
- Graph 7 - (WIP) what is the trend in scores across the subsequent rounds for each member of the Majis team?
- Graph 8 - (still to do) what is the trend in scores across the subsequent rounds for each team?

Files:
- this file is called in dobbelsteen.py (code for generating the scores in each round)

Dependencies:
- dice_scores.xlsx in the working directory

Output files:
- dice_scores_processed.xlsx
- Graph1.png until Graph7.png

"""

import os
import statistics
from statistics import mean
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from tabulate import tabulate

# %%

the_chosen_path = os.path.join(os.path.dirname(__file__))

data     = pd.read_excel(os.path.join(the_chosen_path, "dice_scores.xlsx"))
print(the_chosen_path)
# %%

def print_proportion_table(var_name, dataset):
    print("")
    print(f"Proportions for {var_name}")
    print((dataset[var_name].value_counts().sort_index()) / len(dataset[var_name]))
    print("")


def print_proportion_table_per_group(var_name, group_name, dataset):
    for group_name, group_df in dataset.groupby(group_name):
        print(f"Proportions for {group_name}:")
        print((group_df[var_name].value_counts().sort_index()) / len(group_df[var_name]))
        print()


def get_sumstats(var_name, dataset):
    return f""" Mean {var_name}\t\t{dataset[var_name].mean()}
           Std.dev {var_name}\t\t{dataset[var_name].std()}
           Median {var_name}\t\t{dataset[var_name].median()}
           Min {var_name}\t\t{dataset[var_name].min()}
           Max {var_name}\t\t{dataset[var_name].max()}"""

def tabulate_table_to_word(var_name_1, pretty_name_1, var_name_2, pretty_name_2, printing, var_type_2):
    
    selected_columns = [var_name_1, var_name_2]
    if printing == True:
        print(tabulate(data[selected_columns], headers="keys", tablefmt="pretty"))
    else:
        tabulate(data[selected_columns], headers="keys", tablefmt="pretty")

    doc = Document()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table style
    table.style = "Light Shading Accent 4"
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = f"\n{pretty_name_1}\n"
    hdr_cells[1].text = f"\n{pretty_name_2}\n"
    
    for index, row in data[selected_columns].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = f"\n{row[var_name_1]}\n"
        if var_type_2.lower == "float":
            var_2_formatted = "{:1.f}".format(var_type_2)
            row_cells[1].text = f"\n{row[var_2_formatted]:.1f}\n" if pd.notnull(row[var_2_formatted]) else "\n"
        elif var_type_2.lower == "percentage":
            row_cells[1].text = f"\n{row[var_name_2]:.2%}\n" if pd.notnull(row[var_name_2]) else "\n"
        else:
            row_cells[1].text = f"\n{row[var_name_2]}\n" if pd.notnull(row[var_name_2]) else "\n"
  
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  
                    run.font.name = "Garamond"
    
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "auto")

            for border in ["left", "right"]:
                border_element = OxmlElement(f"w:{border}")
                border_element.set(qn("w:val"), "nil")
                tcBorders.append(border_element)

    for i, row in enumerate(table.rows):
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        else:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = False

    doc.save(f"{the_chosen_path}\{var_name_1}X{var_name_2}.docx")


def proportion_table_to_word(var_name, pretty_name):
    
    variable_proportions = (data[var_name].value_counts().sort_index() / len(data[var_name]))

    doc = Document()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table style: https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
    table.style = "Light Shading Accent 4"
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "\n" + pretty_name + "\n"
    hdr_cells[1].text = "\n" + "Percentages" + "\n"
    
    for variable, proportion in variable_proportions.items():
        row_cells = table.add_row().cells
        row_cells[0].text = f"\n{variable}\n"
        row_cells[1].text = f"\n{proportion:.2%}\n"
  
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            if row_index == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif cell_index == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            for run in paragraph.runs:
                run.font.size = Pt(11)  
                run.font.name = "Garamond"
    
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is None:
                tcBorders = OxmlElement("w:tcBorders")
                tcPr.append(tcBorders)
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "auto")

            for border in ["left", "right"]:
                border_element = OxmlElement(f"w:{border}")
                border_element.set(qn("w:val"), "nil")
                tcBorders.append(border_element)

    for i, row in enumerate(table.rows):
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        else:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = False

    doc.save(f"{the_chosen_path}\proptable_{var_name}.docx")


def save_and_show_plot(filename, transparency):
    plt.savefig(f"{the_chosen_path}/{filename}.png", dpi=500, transparent=transparency, bbox_inches="tight")
    plt.show()


def print_paragraph(text):
    print("\n\n\n\n", 50 * "*", text, 50 * "*", "\n")


def determine_max_y_axis(depvar, dataset):
    max_score = dataset[depvar].max()
    max_y_axis = max_score + max_score * 0.2
    return max_y_axis

def make_graph_7(sourcedata):

    
    palette_for_graph_7 = {
        "Altan": "#d18e75",
        "Goromi": "#176482",
        "Nozomi": "#b6e1f2",
        "Inizio": "#F2C6B6",
        "Cyber": "#663737",
    }
    round_columns = [col for col in data.columns if col.startswith("Round")]

    data_for_graph_7 = sourcedata.melt(id_vars="Participant",
                        value_vars=round_columns,
                        var_name="Round", 
                        value_name="Score")
    
    data_for_graph_7["Round"] = pd.Categorical(data_for_graph_7["Round"], categories=round_columns, ordered=True)

    max_y_axis = determine_max_y_axis("Score", data_for_graph_7)

    plt.figure(figsize=(16, 9))
    for participant in sourcedata["Participant"]:
        participant_data = data_for_graph_7[data_for_graph_7["Participant"] == participant]
        plt.plot(participant_data["Round"], participant_data["Score"], marker="o", label=participant, color=palette_for_graph_7[participant])

    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.title("Scores of Majis Across Different Rounds")
    plt.xlabel("Round", fontsize=32, labelpad=18)
    plt.ylabel("Score", fontsize=32, labelpad=18)
    plt.ylim(-2, max_y_axis)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.legend(title="Participant")
    plt.grid(True)
    save_and_show_plot("Graph7", False)

    return palette_for_graph_7, round_columns, data_for_graph_7, data_for_graph_7["Round"], max_y_axis


# %%

number_of_participants = data.shape[0]
number_of_rounds = data.shape[1] - 3
number_of_team_members = data["Team"].value_counts().sort_index()
largest_team = number_of_team_members.max()

print(f"""Check: 
- you have {number_of_participants} participants;
- you have {number_of_rounds} rounds
- you have {largest_team} members in your largest team.")
""")

# %%

counter = 0
while counter < 1:
    introtext = input("""
                                Hello. Now we're going to analyse the score data and output some graphs!

                                Hit [ENTER] to continue! OR type 'stop' to quit! """)

    if introtext.lower() == "stop":
        break
    elif len(introtext) == 0:
        counter+=1
        pass
    else:
        print_paragraph("You probably made a typo or something went wrong. Try answering again.")
        continue
    
# %%

    participant_names = data["Participant"].tolist()
    participant_colors = {
        participant: sns.color_palette("rocket", n_colors=len(participant_names))[i]
        for i, participant in enumerate(participant_names)
    }
    
    team_names = data["Team"].tolist()
    team_colors = {
        team: sns.color_palette("rocket", n_colors=len(team_names))[i]
        for i, team in enumerate(team_names)
    }

# %%

    round_columns = [f"Round_{round_number}" for round_number in range(1, number_of_rounds + 1)]
    
    data["total_score"] = data[round_columns].sum(axis=1)
    data["average_score"] = data[round_columns].mean(axis=1)

    print_paragraph("Total Score Per Participant:")
    for i in range(1, number_of_participants + 1):
        for index, row in data.iterrows():
            if row["ID"] == i:
                total_score = row["total_score"]
                print(f"{row['Participant']} has a total score of {total_score}")

    print_paragraph("Summary Statistics Total Score Of All Participants:")
    print(get_sumstats("total_score", data))

    print_paragraph("Average Score Per Participant:")
    for i in range(1, number_of_participants + 1):
        for index, row in data.iterrows():
            if row["ID"] == i:
                average_score = row["average_score"]
                print(f"{row['Participant']} has an average score of {average_score}")

    for i in range(1, number_of_rounds + 1):
        globals()[f"highest_round{i}"] = data[f"Round_{i}"].max()
        data[f"top_participant{i}"] = (data[f"Round_{i}"] == globals()[f"highest_round{i}"])

# %%

    rank_mapping = {
        False: 0,
        True: 1}
    for i in range(1, number_of_rounds + 1):
        data[f"top_participant{i}_dummy"] = data[f"top_participant{i}"].map(rank_mapping)
        data[f"top_participant{i}_dummy"] = data[f"top_participant{i}_dummy"].astype(int)

    top_position_columns = [f"top_participant{round_number}_dummy" for round_number in range(1, number_of_rounds + 1)]
    data["total_top_positions"] = data[top_position_columns].sum(axis=1)
    
    print_paragraph("Total Times A Participant Had Top Score:")
    for i in range(1, number_of_participants + 1):
        for index, row in data.iterrows():
            if row["ID"] == i:
                total_top_positions = row["total_top_positions"]
                print(f"{row['Participant']} was in the top at {total_top_positions}")

    print_paragraph("Summary Statistics Top Score Occurences Of All Participants:")
    print(get_sumstats("total_top_positions", data))

    team_top_scores_mean = data.groupby("Team")["total_top_positions"].mean().reset_index()
    team_top_scores_mean.columns = ["Team", "Average No. Of Top Scores For Team"]
    teamdata = pd.merge(data, team_top_scores_mean, on="Team")

# %%

    print_paragraph("Global Info About The Score Dataframe")
    data.info()
    print(f"\n\n {data.head} \n\n")

# %%

    # GRAPH 1
    max_y_axis = determine_max_y_axis("total_score", data)
    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.figure(figsize=(14, 6))
    sns.barplot(data=data, x="Participant", y="total_score", palette=participant_colors, alpha = 0.7, width=1, errorbar = None)
    plt.xlabel("Participant", fontsize=32, labelpad=18)
    plt.ylabel("Total Score", fontsize=32, labelpad=18)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.ylim(0, max_y_axis)
    save_and_show_plot("Graph1", False)

    print_paragraph("TABLE 1 - Total Scores Of Participants")
    tabulate_table_to_word("Participant", "Participant", "total_score", "Total score", True, "string")

    # GRAPH 2
    max_y_axis = determine_max_y_axis("average_score", data)
    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.figure(figsize=(14, 6))
    sns.barplot(data=data, x="Participant", y="average_score", palette=participant_colors, alpha = 0.7, width=1, errorbar = None)
    plt.xlabel("Participant", fontsize=32, labelpad=18)
    plt.ylabel("Average Score", fontsize=32, labelpad=18)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.ylim(0, max_y_axis)
    save_and_show_plot("Graph2", False)

    print_paragraph("TABLE 2 - Average Scores Of Participants")
    tabulate_table_to_word("Participant", "Participant", "average_score", "Average score", True, "float")

    # GRAPH 3
    max_y_axis = determine_max_y_axis("total_top_positions", data)
    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.figure(figsize=(14, 6))
    sns.barplot(data=data, x="Participant", y="total_top_positions", palette=participant_colors, alpha = 0.7, width=1, errorbar = None)
    plt.xlabel("Participant", fontsize=32, labelpad=18)
    plt.ylabel("Total rounds at top", fontsize=32, labelpad=18)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.ylim(0, max_y_axis)
    save_and_show_plot("Graph3", False)

    print_paragraph("TABLE - Total Top Positions for Each Participant")
    print_proportion_table_per_group("total_top_positions", "Participant", data)

    # GRAPH 4
    max_y_axis = determine_max_y_axis("Average No. Of Top Scores For Team", teamdata)
    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.figure(figsize=(14, 9))
    sns.barplot(data=teamdata, x="Team", y="Average No. Of Top Scores For Team", palette=team_colors, alpha = 0.7, width=1, errorbar = None)
    plt.ylabel("Average rounds that team members were at top", fontsize=32, labelpad=18)
    plt.xlabel("Teams", fontsize=32, labelpad=18)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.ylim(0, max_y_axis)
    save_and_show_plot("Graph4", False)

    print_paragraph("TABLE - Average rounds that team members were at top")
    print_proportion_table_per_group("Average No. Of Top Scores For Team", "Team", teamdata)

    # GRAPH 5
    average_total_score = data["total_score"].mean()
    data["final_evaluation"] = data["total_score"].apply(lambda x: 1 if x >= average_total_score else 0)
    above_average_score_data = data[data["final_evaluation"] == 1]
    below_average_score_data = data[data["final_evaluation"] == 0]

    max_y_axis = largest_team + largest_team * 0.2
    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.figure(figsize=(14, 9))
    sns.countplot(data=data, x="Team", hue="final_evaluation", dodge=True, palette={0: "#d18e75", 1: "#176482"}, linewidth=2, alpha=0.7)
    plt.ylabel("Quantity", fontsize=32, labelpad=18)
    plt.xlabel("Teams", fontsize=32, labelpad=18)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.ylim(0, max_y_axis)
    plt.legend(title="Evaluation", loc="upper right", fontsize=22, title_fontsize="24", labels=["Below average", "Above average"])
    save_and_show_plot("Graph5", False)

    print_paragraph("TABLE - Proportion of Scores Above Average per Team")
    print_proportion_table_per_group("final_evaluation", "Team", data)

    # GRAPH 6
    teamdata_rounds = data.melt(id_vars=["Team", "Participant"],
                                value_vars=["Round_1", "Round_2"],
                                var_name="Round",
                                value_name="Score")
    mean_scores_team_round = teamdata_rounds.groupby(["Team", "Round"])["Score"].mean().reset_index()

    print(teamdata_rounds)
    max_y_axis = determine_max_y_axis("Score", mean_scores_team_round)
    sns.set(style="whitegrid", font="Garamond")
    plt.rcParams["font.family"] = "Garamond"
    plt.figure(figsize=(14, 9))
    sns.barplot(data=mean_scores_team_round, x="Team", y="Score", hue="Round", palette={"Round_1": "#d18e75", "Round_2": "#176482"}, alpha = 0.7, linewidth=2, errorbar = None)
    plt.ylabel("Average score", fontsize=32, labelpad=18)
    plt.xlabel("Teams", fontsize=32, labelpad=18)
    plt.xticks(fontsize=24, rotation=90)
    plt.yticks(fontsize=24)
    plt.ylim(0, max_y_axis)
    save_and_show_plot("Graph6", False)

    print_paragraph("TABLE - Average score for teams in Round 1")
    print(data.groupby(["Team"])["Round_1"].mean().reset_index)

  
    # GRAPH 7 (if user has chosen default participants)

    if "Majis" in data["Team"].values:
        selection_for_graph_7 = data[data["Team"] == "Majis"]
        make_graph_7(selection_for_graph_7)
    elif "Pocky Lovers" in data["Team"].values:
        make_graph_7(data)
    


    #data_for_graph_7 = selection_for_graph_7.melt(id_vars="Participant",
     #                            value_vars=round_columns,
      #                           var_name="Round", 
       #                          value_name="Score")
    



    #sns.set(style="whitegrid", font="Garamond")
    #plt.rcParams["font.family"] = "Garamond"
    #plt.figure(figsize=(14, 9))

    #sns.barplot(data=data, x="Team", y="total_top_positions", hue="Team", ci=None, palette={"Majis": "#d18e75", "Dragon Team": "#176482", "Baddies": "#89c6b6", "Japanese Plushies": "pink", "Hydrogenpink Plushies": "#F2C6B6", "Phantom Thieves": "orange", "House Lady": "#60272b"}, alpha=0.7)
    #plt.ylabel("Total rounds at top", fontsize=32, labelpad=18)
    #plt.xlabel("Teams", fontsize=32, labelpad=18)
    #plt.xticks(fontsize=24, rotation=90)
    #plt.yticks(fontsize=24)
    #plt.ylim(0, 5)
    #plt.legend(title="Team", loc="upper right", fontsize=22, title_fontsize="24", labels=["Majis", "Dragon Team", "Baddies", "Japanese Plushies", "Hydrogenpink Plushies", "Phantom Thieves", "House Lady"], ncol=3)
    #save_and_show_plot("scores4", False)

    #print_proportion_table_per_group("total_top_positions", "Team", data)

# %%

    data.to_excel(os.path.join(the_chosen_path, "dice_scores_processed.xlsx"), index=False)

# %%

    input("Press key to end ")

    # Actiepunten:
    # Maak de code mooier met #%# of zoiets en docstring
    # Maak de grafieken nog mooier! Legenda wel of niet. Kleuren goed?
    # De twee andere grafieken ook nog doen (Graph 6 en 7)
    # Kijken waar meer functies kunnen worden gedefinieerd.
    # Verder met de tabellen. Tabellen bepalen: wat wil je in tabellen in de report
    # Report maken met tabellen, grafieken en tekst. Allemaal geautomatiseerd!
    # Kijk of je code HELEMAAL object oriented kan! En efficienter.
